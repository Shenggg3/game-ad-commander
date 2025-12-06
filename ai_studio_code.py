import streamlit as st
import google.generativeai as genai
import openai
import urllib.parse
import random
from PIL import Image
import io
import base64
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 1. 頁面配置與 CSS
# ==========================================
st.set_page_config(
    page_title="全球遊戲廣告素材指揮官 (V15.0 雙引擎旗艦版)",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    /* 全局設定 */
    .stApp { background-color: #0F172A; color: #E2E8F0; }
    
    /* 標題特效 */
    .title-text { 
        color: #60A5FA; 
        text-align: center; 
        font-weight: 800; 
        letter-spacing: 2px; 
        font-size: 2.5em; 
        text-shadow: 0 0 20px rgba(96, 165, 250, 0.4); 
    }
    
    /* 步驟標題 */
    .step-header {
        background: linear-gradient(90deg, #1e293b 0%, #0f172a 100%);
        padding: 15px;
        border-radius: 8px;
        border-left: 5px solid #60A5FA;
        font-size: 1.2em;
        font-weight: bold;
        color: #60A5FA;
        margin-top: 20px;
        margin-bottom: 15px;
    }

    /* 區塊樣式 */
    .box-style { background-color: #1E293B; padding: 20px; border-radius: 10px; border: 1px solid #334155; }
    .scene-card { background-color: #1E293B; border: 1px solid #475569; border-radius: 12px; padding: 20px; margin-bottom: 25px; border-left: 6px solid #38BDF8; box-shadow: 0 4px 10px rgba(0,0,0,0.3); }
    .video-prompt-box { background-color: #020617; border: 1px dashed #4ADE80; padding: 12px; border-radius: 6px; font-family: 'Courier New', monospace; color: #4ADE80; font-size: 0.85em; margin-top: 10px; }
    
    /* 聲音標籤 */
    .audio-vo { color: #FACC15; font-weight: bold; }
    .audio-dia { color: #E879F9; font-weight: bold; }
    .audio-sfx { color: #F87171; font-weight: bold; font-size: 0.9em; }
</style>
""", unsafe_allow_html=True)

# Session State 初始化
if 'game_analysis_result' not in st.session_state: st.session_state.game_analysis_result = {}
if 'current_step' not in st.session_state: st.session_state.current_step = 1
if 'final_script_data' not in st.session_state: st.session_state.final_script_data = None 

# ==========================================
# 2. 核心 AI 函數 (雙引擎切換)
# ==========================================
def encode_image_to_base64(image):
    """將 PIL Image 轉換為 base64 字串 (給 OpenAI 用)"""
    buffered = io.BytesIO()
    image.save(buffered, format="JPEG")
    return base64.b64encode(buffered.getvalue()).decode('utf-8')

def call_ai_engine(provider, api_key, model_name, system_prompt, user_prompt, image=None):
    """統一呼叫介面：處理 Google 與 OpenAI 的差異"""
    try:
        if provider == "Google Gemini":
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            
            # Gemini 的輸入列表
            inputs = [system_prompt + "\n\n" + user_prompt]
            if image:
                inputs.append(image)
                
            response = model.generate_content(inputs)
            return response.text

        elif provider == "OpenAI":
            client = openai.OpenAI(api_key=api_key)
            
            messages = [
                {"role": "system", "content": system_prompt},
            ]
            
            user_content = []
            user_content.append({"type": "text", "text": user_prompt})
            
            if image:
                base64_image = encode_image_to_base64(image)
                user_content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{base64_image}"
                    }
                })
            
            messages.append({"role": "user", "content": user_content})
            
            response = client.chat.completions.create(
                model=model_name,
                messages=messages
            )
            return response.choices[0].message.content
            
    except Exception as e:
        raise Exception(f"{provider} API Error: {str(e)}")

# ==========================================
# 3. 輔助函數：Word 生成
# ==========================================
def generate_docx(game_name, strategy, scenes_data):
    doc = Document()
    heading = doc.add_heading(f'廣告腳本企劃書: {game_name}', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('🧠 行銷心理戰略', level=1)
    p_strat = doc.add_paragraph(strategy)
    p_strat.paragraph_format.space_after = Pt(12)
    
    doc.add_heading('📋 分鏡詳細腳本', level=1)
    
    for i, scene in enumerate(scenes_data):
        doc.add_heading(f'Scene {i+1} ({scene.get("Time", "N/A")})', level=2)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        
        run = p.add_run('🎥 畫面: '); run.bold = True; run.font.color.rgb = RGBColor(0, 0, 0)
        p.add_run(f"{scene.get('Visual', '')}\n")
        
        run = p.add_run('📝 壓字: '); run.bold = True
        p.add_run(f"{scene.get('Text', '')}\n")
        
        if scene.get('Voiceover') not in ["None", "無"]:
            run = p.add_run('🗣️ 旁白: '); run.bold = True; run.font.color.rgb = RGBColor(0, 112, 192)
            p.add_run(f"{scene.get('Voiceover', '')}\n")
            
        if scene.get('Dialogue') not in ["None", "無"]:
            run = p.add_run('💬 對話: '); run.bold = True; run.font.color.rgb = RGBColor(112, 48, 160)
            p.add_run(f"{scene.get('Dialogue', '')}\n")
            
        run = p.add_run('🔊 音效: '); run.bold = True; run.font.color.rgb = RGBColor(192, 0, 0)
        p.add_run(f"{scene.get('SFX', '')}\n")
        
        p_prompt = doc.add_paragraph()
        run_label = p_prompt.add_run('Video AI Prompt: '); run_label.bold = True; run_label.font.size = Pt(9)
        run_text = p_prompt.add_run(f"{scene.get('Video Prompt', '')}"); run_text.italic = True; run_text.font.size = Pt(9); run_text.font.color.rgb = RGBColor(80, 80, 80)
        
        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 4. 側邊欄：雙引擎設定
# ==========================================
with st.sidebar:
    st.title("⚙️ 引擎設定")
    
    # 選擇供應商
    ai_provider = st.selectbox("🤖 選擇 AI 供應商", ["Google Gemini", "OpenAI"])
    
    api_key = st.text_input(f"🔑 輸入 {ai_provider} API Key", type="password")
    
    # 根據供應商選擇模型
    model_name = ""
    if ai_provider == "Google Gemini":
        model_name = "gemini-1.5-flash" # 預設，可改
        st.caption("預設使用: gemini-1.5-flash (速度快、免費額度高)")
    else:
        model_name = st.selectbox("🧠 選擇 OpenAI 模型", ["gpt-4o", "gpt-4-turbo", "gpt-3.5-turbo"])
        st.caption("推薦使用: gpt-4o (視覺能力最強)")
        
    st.divider()
    
    if st.button("🔗 測試連線"):
        if not api_key:
            st.error("請輸入 API Key")
        else:
            try:
                # 簡單測試呼叫
                with st.spinner(f"正在測試 {ai_provider} 連線..."):
                    test_res = call_ai_engine(ai_provider, api_key, model_name, "System", "Hello, say hi!", None)
                    st.success(f"✅ 連線成功！回應: {test_res}")
            except Exception as e:
                st.error(f"連線失敗: {e}")

# ==========================================
# 5. 主畫面
# ==========================================
st.markdown("<h1 class='title-text'>🌍 全球遊戲廣告素材指揮官</h1>", unsafe_allow_html=True)
st.markdown(f"<p style='text-align: center; color: #94A3B8;'>V15.0 雙引擎版 • 當前引擎: <span style='color:#60A5FA; font-weight:bold;'>{ai_provider} ({model_name})</span></p>", unsafe_allow_html=True)

# ------------------------------------------
# STEP 1: 遊戲調研 (Research)
# ------------------------------------------
st.markdown('<div class="step-header">STEP 1: 遊戲視覺與戰略調研</div>', unsafe_allow_html=True)

c_g1, c_g2 = st.columns([2, 1])
with c_g1: game_name_input = st.text_input("🎮 遊戲名稱", placeholder="Ex: 絕區零")
with c_g2: platform_input = st.selectbox("🕹️ 遊戲平台", ["手機遊戲", "PC/Steam", "主機", "網頁遊戲"])

uploaded_game_img = st.file_uploader("📸 (選填) 上傳截圖，讓 AI 用眼睛看", type=["jpg", "png", "jpeg"])

if st.button("👁️ 啟動視覺調研"):
    if not api_key or not game_name_input:
        st.warning("請先輸入 API Key 並填寫遊戲名稱")
    else:
        with st.spinner(f"[{ai_provider}] 正在分析《{game_name_input}》..."):
            system_prompt = "You are a Senior Game Producer."
            user_prompt = f"""
            Analyze game "{game_name_input}" on "{platform_input}".
            
            **Task:**
            1. Identify Genre & Core Loop.
            2. Identify 3 USP (Unique Selling Points).
            3. **Visual Analysis:** Describe the art style, color palette, UI style, and character proportions in detail.
            
            **If image is provided:** Use it to describe the Visual Style accurately.
            
            Output strictly in Traditional Chinese:
            Genre: [類型]
            Core Loop: [核心玩法]
            USP: [3個賣點]
            Visual Style: [美術風格 - 詳細描述]
            """
            
            img_obj = Image.open(uploaded_game_img) if uploaded_game_img else None
            
            try:
                res_text = call_ai_engine(ai_provider, api_key, model_name, system_prompt, user_prompt, img_obj)
                
                st.session_state.game_analysis_result = {
                    "name": game_name_input,
                    "platform": platform_input,
                    "raw_analysis": res_text
                }
                st.session_state.current_step = 2
            except Exception as e:
                st.error(f"調研失敗: {e}")

if st.session_state.current_step >= 2:
    st.markdown('<div class="box-style">', unsafe_allow_html=True)
    st.info("👇 AI 的遊戲戰略認知")
    game_profile_user_edit = st.text_area(
        "📝 遊戲戰略檔案 (可修正):",
        value=st.session_state.game_analysis_result.get("raw_analysis", ""),
        height=150
    )
    st.markdown('</div>', unsafe_allow_html=True)

# ------------------------------------------
# STEP 2: 參數與生成 (Generation)
# ------------------------------------------
if st.session_state.current_step >= 2:
    st.markdown('<div class="step-header">STEP 2: 創意生成</div>', unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1: target_region = st.selectbox("🌐 投放地區", ["台灣 (繁中)", "日本 (日文)", "美國 (英文)", "韓國 (韓文)", "中國大陸 (簡中)", "東南亞"])
    with c2: duration = st.select_slider("⏱️ 廣告時長", options=[15, 30, 45, 60], value=30)

    c3, c4 = st.columns(2)
    with c3:
        tone_sel = st.selectbox("🎭 影片風格", ["搞笑諧音", "熱血中二", "懸疑驚悚", "感人共鳴", "專業硬核", "✨ 自定義"])
        ad_tone = st.text_input("✍️ 自定義風格", placeholder="Ex: 王家衛風") if tone_sel == "✨ 自定義" else tone_sel
    with c4:
        fmt_sel = st.selectbox("📢 腳本形式", ["戰力飆升", "失敗挑戰", "CG 動畫大片", "實機試玩", "真人情境劇", "✨ 自定義"])
        ad_format = st.text_input("✍️ 自定義形式", placeholder="Ex: 靈魂互換") if fmt_sel == "✨ 自定義" else fmt_sel

    st.markdown("<b>🎯 受眾與情境</b>", unsafe_allow_html=True)
    col_demo, col_context = st.columns([3, 2])
    with col_demo:
        t1, t2, t3 = st.columns(3)
        with t1: ta_gender = st.selectbox("👤 性別", ["不限", "男性", "女性"])
        with t2: ta_age = st.slider("🎂 年齡", 12, 60, (25, 35))
        with t3: ta_identity = st.text_input("💼 身分", value="上班族")
    with col_context:
        t4, t5 = st.columns(2)
        with t4: ta_time = st.selectbox("⏰ 投放時段", ["通勤/上學 (早上)", "午休時間 (中午)", "下班/放學 (晚上)", "深夜時段 (半夜)", "全天候"])
        with t5: ta_holiday = st.text_input("🎉 節慶", value="平日")

    with st.expander("📝 導演筆記 (補充指令)"):
        custom_req = st.text_area("特殊需求...", placeholder="Ex: 結局要有反轉")

    if st.button("🚀 啟動引擎生成腳本"):
        if not api_key:
             st.error("API Key 遺失，請重新輸入")
        else:
            with st.spinner(f"[{ai_provider}] 正在進行戰略運算與分鏡撰寫..."):
                system_prompt = "You are a World-Class Game Creative Director."
                user_prompt = f"""
                **INPUTS:**
                - Game Profile: {game_profile_user_edit}
                - Region: {target_region}
                - Duration: {duration}s
                - Tone: {ad_tone}
                - Format: {ad_format}
                - Audience: {ta_identity} ({ta_gender}, Age {ta_age[0]}-{ta_age[1]})
                - Context: Time: {ta_time}, Holiday: {ta_holiday}
                - User Note: {custom_req}
                
                **TASK:**
                1. **Psych Strategy:** Map Game USP to User Pain Points.
                2. **Script:** Scene-by-scene breakdown.
                   - Voiceover/Dialogue: Native Language of {target_region}.
                   - Visuals: Traditional Chinese.
                   - Audio: Separate Voiceover/Dialogue/SFX.
                3. **Video Prompt:** English for Sora/Veo3 (Focus on Motion & Physics).
                
                **OUTPUT FORMAT (Separator '|||'):**
                
                [STRATEGY]
                心理戰略: [Analysis]
                |||
                Scene 1
                Time: [Seconds]
                Visual: [Desc]
                Voiceover: [Script]
                Dialogue: [Script]
                SFX: [Desc]
                Text: [Overlay]
                Video Prompt: [English Prompt]
                |||
                (Repeat)
                """

                try:
                    res_text = call_ai_engine(ai_provider, api_key, model_name, system_prompt, user_prompt, None)
                    
                    if "[STRATEGY]" in res_text:
                        parts = res_text.split("|||")
                        strategy = parts[0].replace("[STRATEGY]", "").strip()
                        scenes_raw = parts[1:]
                    else:
                        strategy = "無策略分析"
                        scenes_raw = res_text.split("|||")
                    
                    parsed_scenes = []
                    for scene in scenes_raw:
                        if len(scene.strip()) < 10: continue
                        lines = scene.strip().split('\n')
                        data = {"Time": "N/A", "Visual": "無", "Voiceover": "無", "Dialogue": "無", "SFX": "無", "Text": "無", "Video Prompt": ""}
                        for line in lines:
                            for k in data.keys():
                                if f"{k}:" in line: data[k] = line.split(":", 1)[1].strip()
                        parsed_scenes.append(data)
                    
                    st.session_state.final_script_data = {
                        "strategy": strategy,
                        "scenes": parsed_scenes,
                        "game_name": game_name_input
                    }
                    st.success("生成完成！")

                except Exception as e:
                    st.error(f"生成錯誤: {e}")

# ------------------------------------------
# STEP 3: 顯示與下載 (Export)
# ------------------------------------------
if st.session_state.final_script_data:
    data = st.session_state.final_script_data
    
    st.markdown(f'<div class="box-style" style="border-left:4px solid #38BDF8;"><h3>🧠 策略</h3>{data["strategy"]}</div><br>', unsafe_allow_html=True)
    
    for i, scene in enumerate(data['scenes']):
        with st.container():
            c_txt, c_img = st.columns([3, 2])
            with c_txt:
                audio_html = ""
                if scene['Voiceover'] not in ["None", "無"]: audio_html += f'<span class="audio-vo">🗣️ 旁白:</span> {scene["Voiceover"]}<br>'
                if scene['Dialogue'] not in ["None", "無"]: audio_html += f'<span class="audio-dia">💬 對話:</span> {scene["Dialogue"]}<br>'
                
                st.markdown(f"""
                <div class="scene-card">
                    <span style="background:#38BDF8; color:#000; padding:2px 6px; border-radius:4px;">Scene {i+1} | {scene['Time']}</span>
                    <br><br><b>🎥 畫面:</b> {scene['Visual']}<br>
                    {audio_html}
                    <span class="audio-sfx">🔊 音效:</span> {scene['SFX']}
                </div>
                """, unsafe_allow_html=True)
                st.markdown(f'<div class="video-prompt-box">{scene["Video Prompt"]}</div>', unsafe_allow_html=True)
            
            with c_img:
                if scene['Video Prompt']:
                    w, h, ratio = (576, 1024, "9:16") if "手機" in st.session_state.game_analysis_result.get('platform', '') else (1024, 576, "16:9")
                    clean_p = urllib.parse.quote(f"{scene['Video Prompt']}, {game_name_input} style, best quality")
                    seed = random.randint(0, 9999)
                    url = f"https://image.pollinations.ai/prompt/{clean_p}?width={w}&height={h}&seed={seed}&nologo=true&model=flux"
                    st.image(url, caption=f"視覺示意 ({ratio})", use_container_width=True)

    st.markdown('<div class="step-header">STEP 3: 商業文件匯出 (Word)</div>', unsafe_allow_html=True)
    
    docx_file = generate_docx(data['game_name'], data['strategy'], data['scenes'])
    
    col_dl_btn, col_dl_info = st.columns([1, 2])
    with col_dl_btn:
        st.download_button(
            label="📄 下載 Word 腳本 (.docx)",
            data=docx_file,
            file_name=f"{data['game_name']}_廣告腳本.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col_dl_info:
        st.info("💡 支援 Google Gemini 與 OpenAI GPT-4o 雙引擎輸出。")